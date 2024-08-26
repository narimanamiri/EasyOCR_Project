
# Importing the different libraries
import cv2 
import numpy as np
import easyocr
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Defining the path of images
img0_path = './images/text0.jpg'
img1_path = './images/text1.jpg'
img2_path = './images/text2.jpg'
img3_path = './images/text3.jpg'
img4_path = './images/text4.jpg'
img5_path = './images/text5.jpg'
img6_path = './images/text6.jpg'
img7_path = './images/text7.jpg'
img8_path = './images/text8.jpg'
img9_path = './images/text9.jpg'

# Extracting the path of images one by one in the form of list
img_path = "./images/"
create_path = lambda f : os.path.join(img_path, f)
test_image_files = os.listdir(img_path)

for f in test_image_files:
    print(f)

# loading the path of image by passing the position of the image
img = test_image_files[8]
path = create_path(img)

# Recognizing the text from images
def recognize_text(img_path):
    ''' loads an image and recognizes text. '''
    
    reader = easyocr.Reader(['en', 'hi'])        # For English - 'en'  And,  For Hindi - 'hi'
    return reader.readtext(img_path)

result = recognize_text(path)
print(result)

# Loading the image
img_1 = cv2.imread(path)
img_1 = cv2.cvtColor(img_1, cv2.COLOR_BGR2RGB)
plt.imshow(img_1)

# Process of Extraction
def overlay_ocr_text(img_path, save_name):
    ''' loads an image, recognize text, and overlays the text on the image. '''
    
    # Load images
    img = cv2.imread(img_path)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    
    dpi = 100
    fig_width, fig_height = int(img.shape[0]/dpi), int(img.shape[1]/dpi)
    plt.figure()
    f, axarr = plt.subplots(1,2, figsize=(fig_width, fig_height))
    axarr[0].imshow(img)
    
    # Recognise text
    result = recognize_text(img_path)
    
    # If OCR prob is over 0.2, overlay bounding box and text
    for(bbox, text, prob) in result:
        if prob >= 0.2:
            print(f'Detected text: {text} (Probability: {prob:.2f})')
            
            # get top-left and bottom-right bbox vertices
            (top_left, top_right, bottom_right, bottom_left) = bbox
            top_left = (int(top_left[0]), int(top_left[1]))
            bottom_right = (int(bottom_right[0]), int(bottom_right[1]))
            
            # create a rectangle for bbox display
            cv2.rectangle(img=img, pt1=top_left, pt2=bottom_right, color=(255,0,0), thickness=10)
            
            # put recognized text
            cv2.putText(img=img, text=text, org=(top_left[0], top_left[1] - 10), fontFace=cv2.FONT_HERSHEY_SIMPLEX, fontScale=1, color=(255,0,0), thickness=8)
            
    # show and save image
    axarr[1].imshow(img)
    plt.savefig(f'./Output/{save_name}_overlay.jpg', bbox_inches='tight')

overlay_ocr_text(path, 'Text3_Output')

# Final Output(text)
def ocr_text(img_path):
    result = recognize_text(img_path)
    
    # Create a Word document
    doc = Document()
    
    # If OCR prob is over 0.2, overlay text
    for(bbox, text, prob) in result:
        if prob >= 0.2:
            print(f'{text}')
            # Check if the text is part of a table
            if 'table' in text.lower():
                # Create a table in the document
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Column 1'
                hdr_cells[1].text = 'Column 2'
                row_cells = table.add_row().cells
                row_cells[0].text = text
                row_cells[1].text = 'Data'
            else:
                paragraph = doc.add_paragraph(text)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = paragraph.runs[0]
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    
    # Save the document
    doc.save('extracted_text.docx')
            
# Extracted Text
ocr_text(path)
